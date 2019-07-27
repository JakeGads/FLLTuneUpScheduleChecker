from tkinter import filedialog
import os
from datetime import time
# User defined
import xlrd

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH # this does exist even if VS tells you otherwise

def reversColConvert(x):
    alpha = ['a','b','c','d','e','f','g','h','i','j','k','l','m','n','o','p','q','r','s','t','u','v','w','x','y','z']
    length = len(alpha)
    if length >= 2:
        #creates the dual letter colums
        for i in range (0, length):
            for h in range (0, length):
                alpha.append((alpha[i] + alpha[h]))
        if length >= 3:
            # creates the tri colums
            for i in range (0, length): # for(int i = 0; i < length; i++)
                for h in range (0, length):
                    for j in range (0, length):
                        alpha.append((alpha[i] + alpha[h] + alpha[j]))            

    return alpha[x]

def verticalCheck(file = None):
    try:
        os.remove('verticalCheck.txt')
    except:
        None
    
    if file is None:
        book = xlrd.open_workbook(filedialog.askopenfilename(title = "Select File",filetypes = (("xlsx files","*.xlsx"),("xls files","*.xls"),("all files","*.*"))))
    else:
        book = xlrd.open_workbook(file)

    sheet = book.sheet_by_index(0)

    nrows = sheet.nrows - 1
    ncols = sheet.ncols - 1

    check = True

    for col in range(ncols):
        if col in [0,1,2,3]:
            continue

        data = []
        for row in range(nrows):
            value = sheet.cell(rowx=row, colx=col).value

            if value == '' or value == 'Judge’s Break' or value == 'Coach Meeting' or value == 'Opening Ceremony' or value == 'Line Dancing':
                continue

            if value in data:
                log = open('verticalCheck.txt', 'a+')
                log.write('col:{col}\trow:{row}\trepeat_value:{value}\n'.format(col= reversColConvert(col), row=row + 1, value=value))
                log.close()
                check = False
            else:
                data.append(value)

    if check:
        print('Vertical Checks passed')
    else:
        print('1 or more Vertical Checks failed see more info @ verticalCheck.txt')
    
    return check

def roomCheck(file=None):
    try:
        os.remove('roomCheck.txt')
    except:
        None
    
    if file is None:
        book = xlrd.open_workbook(filedialog.askopenfilename(title = "Select File",filetypes = (("xlsx files","*.xlsx"),("xls files","*.xls"),("all files","*.*"))))
    else:
        book = xlrd.open_workbook(file)

    sheet = book.sheet_by_index(0)
    check = True
    roomCol = 0

    

    try:
        for row in range(sheet.nrows):
            firstTeam = sheet.cell(rowx=row, colx=roomCol).value.replace('-ii','').replace('-i','')
            secondTeam = sheet.cell(rowx=row + 1, colx=roomCol).value.replace('-ii','').replace('-i','')

            if firstTeam != secondTeam:
                continue

            for col in range(sheet.ncols):
                first_cell = sheet.cell(rowx=row, colx=col).value
                second_cell = sheet.cell(rowx=row + 1, colx=col).value

                if first_cell == 'Judge’s Break' or first_cell == 'Coach Meeting' or first_cell == 'Opening Ceremony' or first_cell == 'Line Dancing':
                    continue

                if first_cell == '' and second_cell == '':
                    continue
                if (first_cell == '' and 'field' in second_cell) or ('field' in first_cell and second_cell == ''):
                    continue

                if 'judge' in first_cell.lower() or 'judge' in second_cell.lower():
                    if 'field' in first_cell.lower() or 'field' in second_cell.lower():
                        None
                    else:
                        check = False
                        log = open('roomCheck.txt', 'a+')
                        log.write('col:{col}\trow:{row}\n'.format(col=reversColConvert(col), row=row + 1))
                        log.close()
                            
    
                    
    except:
        None

    if check:
        print('RoomCheck has been cleared')
    else:
        print('1 or more room failed please check @ roomCheck.txt for more details')

    return check

def generateTeamDocs(file = None):
    class TeamEvent():
        def __init__(self, title, time, session):
            self.title = title
            self.session = session
            self.time = time
        def __str__(self):
            return '{time}:{session}\t{title}'.format(time=self.time, session=self.session, title=self.title)

    def time_convert(x):
        x = int(x * 24 * 3600) # convert to number of seconds
        hour = time(x//3600, (x%3600)//60).hour
        if hour > 12:
            hour -= 12
        minute = time(x//3600, (x%3600)//60).minute
        if minute < 10:
            minute = '0' + str(minute)
        return str(hour) + ':' + str(minute) 

    def main(file = None):
        try:
            os.mkdir('teamSchedule')
        except:
            None

        if file is None:
            book = xlrd.open_workbook(filedialog.askopenfilename(title = "Select File",filetypes = (("xlsx files","*.xlsx"),("xls files","*.xls"),("all files","*.*"))))
        else:
            book = xlrd.open_workbook(file)

        sheet = book.sheet_by_index(0)
        multiDoc = Document()
        for row in range(sheet.nrows):
            singleDoc = Document()
            if row is 0 or row is 1:
                continue
            
            sched = []
            room = ''
            number = ''
            name = ''
            for col in range(sheet.ncols):
                if col is 0:
                    room = sheet.cell(rowx=row,colx=col).value.replace('-ii','').replace('-i','')
                    continue
                if col is 1:
                    number = sheet.cell(rowx=row,colx=col).value
                    continue
                if col is 2:
                    name = sheet.cell(rowx=row,colx=col).value
                    continue
                if col is 3:
                    continue

                if sheet.cell_type(row, col) != xlrd.empty_cell and sheet.cell_value(row,col) != '':
                    sched.append(TeamEvent(sheet.cell(rowx=row,colx=col).value, time_convert(sheet.cell(rowx=1,colx=col).value), sheet.cell(rowx=0,colx=col).value))

            # print('\n\n{room}\t{number}\t{name}'.format(room=room, number=number, name=name))
            
            # for i in sched:
            #    print('\t\t\t'+str(i))
            try: 
                singleDoc.add_heading('{number}\t\t{name}'.format(number=int(number), name=name))
                multiDoc.add_heading('{number}\t\t{name}'.format(number=int(number), name=name))
            except:
                break

            try:
                singleDoc.add_paragraph('Room {room}'.format(room=int(room)))
                multiDoc.add_paragraph('Room {room}'.format(room=int(room)))
            except:
                singleDoc.add_paragraph('Room {room}'.format(room=room))
                multiDoc.add_paragraph('Room {room}'.format(room=room))

            singleTable = singleDoc.add_table(rows=1, cols=3)
            singleHeader = singleTable.rows[0].cells
            singleHeader[0].text = 'Time'
            singleHeader[1].text = 'Session'
            singleHeader[2].text = 'Event'

            for i in sched:
                tbRow = singleTable.add_row().cells
                tbRow[0].text = i.time
                tbRow[1].text = i.session
                tbRow[2].text = i.title


            multiTable = multiDoc.add_table(rows=1, cols=3)
            multiHeader = multiTable.rows[0].cells
            multiHeader[0].text = 'Time'
            multiHeader[1].text = 'Session'
            multiHeader[2].text = 'Event'

            for i in sched:
                tbRow = multiTable.add_row().cells
                tbRow[0].text = i.time
                tbRow[1].text = i.session
                tbRow[2].text = i.title   # VS will through an error this is actual true

            singleCentered = singleDoc.add_paragraph('\n\nAll times are approximate please refer to the session markers')
            singleCentered.alignment = WD_ALIGN_PARAGRAPH.CENTER
            multiCentered = multiDoc.add_paragraph('\n\nAll times are approximate please refer to the session markers')
            multiCentered.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                singleDoc.save('teamSchedule//{num}_{name}_schedule.docx'.format(num=int(number), name=name))
            except:
                singleDoc.save('teamSchedule//{num}_{name}_schedule.docx'.format(num=number, name=name))
        multiDoc.save('teamSchedule//0_Full_List.docx')
        
    main(file=file)

def generateJudgeDocs(file=None):
    class JudgeEvent():
        def __init__(self, teamNumber, teamName, room, time, session):
            self.teamNumber = teamNumber
            self.teamName = teamName
            self.room = room
            self.time = time
            self.session = session

    def time_convert(x):
        x = int(x * 24 * 3600) # convert to number of seconds
        hour = time(x//3600, (x%3600)//60).hour
        if hour > 12:
            hour -= 12
        minute = time(x//3600, (x%3600)//60).minute
        if minute < 10:
            minute = '0' + str(minute)
        return str(hour) + ':' + str(minute) 

    def findJudges(sheet):
        hits = []
    
        
        for col in range(sheet.ncols-1):
            for row in range(sheet.nrows-1):
                val = sheet.cell(colx=col, rowx=row).value
                try:
                    if 'judge' in val.lower() and val not in hits:
                        hits.append(val)
                except:
                    None
        
        return hits


    def main(file = None):
        try:
            os.mkdir('judgeSchedule')
        except:
            None

        try:
            sheet = xlrd.open_workbook(file).sheet_by_index(0)
        except:
            sheet= xlrd.open_workbook(filedialog.askopenfilename(title = "Select File",filetypes = (("xlsx files","*.xlsx"),("xls files","*.xls"),("all files","*.*")))).sheet_by_index(0)

        judgeGroups = findJudges(sheet)

        session_row = 0
        time_row = 1 

        room_col = 0
        team_num_col = 1
        team_name_col = 2


        multiDoc = Document()
        for group in judgeGroups:
            judgeSchedule = []
            singleDoc = Document()
            for col in range(sheet.ncols-1):
                for row in range(sheet.nrows-1):
                    val = sheet.cell(colx=col, rowx=row).value

                    if val == group:
                        # teamNumber, teamName, room, time, session
                        judgeSchedule.append(
                            JudgeEvent(
                                sheet.cell(colx=team_num_col, rowx=row).value, 
                                sheet.cell(colx=team_name_col, rowx=row).value, 
                                sheet.cell(colx=room_col, rowx=row).value,
                                time_convert(sheet.cell(colx=col, rowx= time_row).value),
                                sheet.cell(colx=col, rowx= session_row).value
                            )
                        )
            singleDoc.add_heading(group)
            multiDoc.add_heading(group)

            table = singleDoc.add_table(rows=1, cols=5)
            tbHead = table.rows[0].cells
            tbHead[0].text = 'Time'
            tbHead[1].text = 'Session'
            tbHead[2].text = 'Room'
            tbHead[3].text = 'Team Number'
            tbHead[4].text = 'Team Name'

            for i in judgeSchedule:
                tbRow = table.add_row().cells
                tbRow[0].text = i.time
                tbRow[1].text = i.session
                tbRow[2].text = i.room
                tbRow[3].text = str(int(i.teamNumber))
                tbRow[4].text = i.teamName

            table = multiDoc.add_table(rows=1, cols=5)
            tbHead = table.rows[0].cells
            tbHead[0].text = 'Time'
            tbHead[1].text = 'Session'
            tbHead[2].text = 'Room'
            tbHead[3].text = 'Team Number'
            tbHead[4].text = 'Team Name'

            for i in judgeSchedule:
                tbRow = table.add_row().cells
                tbRow[0].text = i.time
                tbRow[1].text = i.session
                tbRow[2].text = i.room
                tbRow[3].text = str(int(i.teamNumber))
                tbRow[4].text = i.teamName
            
            multiDoc.add_page_break()
            singleDoc.save('judgeSchedule//{group}.docx'.format(group=group))
        multiDoc.save('judgeSchedule//Full_Schedule.docx')

            
            
        

    main(file=file)

def pdfRender(folder):
    from subprocess import call
    og = os.getcwd()

    call('cd {loc}'.format(loc = folder))

    for i in os.listdir():
        # convert


    call('cd {og}'.format(og=og))

if __name__ == "__main__":
    try:
        file = 'FLL Schedule.xlsx'
    except:
        file = filedialog.askopenfilename(title = "Select File",filetypes = (("xlsx files","*.xlsx"),("xls files","*.xls"),("all files","*.*")))


    if verticalCheck(file=file) and roomCheck(file=file):
        print('generating team schedules')
        generateTeamDocs(file=file) 
        print('generating judge schedules')
        generateJudgeDocs(file=file)

        if os.name == 'Windows':
            print('rendering docs as pdf')
            pdfRender('judgeSchedule')
            pdfRender('teamSchedule')
        else:
            print('PDF rendering is only avalible on windows')
        
        